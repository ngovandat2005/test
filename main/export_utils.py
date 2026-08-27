import os
from datetime import datetime
from django.conf import settings
from django.utils import timezone
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def get_export_dir():
    export_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
    os.makedirs(export_dir, exist_ok=True)
    return export_dir


def export_consultations_to_excel(queryset, workbook=None):
    """
    Xuất danh sách đăng ký tư vấn sang workbook openpyxl được định dạng đẹp mắt
    """
    if workbook is None:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Đăng Ký Tư Vấn"
    else:
        wb = workbook
        ws = wb.active

    # Cấu hình font & style
    header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    header_fill = Font(name="Arial", size=11, bold=True)
    orange_fill = PatternFill(start_color="E65A1E", end_color="E65A1E", fill_type="solid")
    title_font = Font(name="Arial", size=15, bold=True, color="E65A1E")
    meta_font = Font(name="Arial", size=10, italic=True, color="666666")
    data_font = Font(name="Arial", size=10)
    data_bold = Font(name="Arial", size=10, bold=True)
    
    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )

    # Tiêu đề bảng tính
    ws.merge_cells('A1:G1')
    ws['A1'] = "DANH SÁCH KHÁCH HÀNG ĐĂNG KÝ TƯ VẤN - SHK MORTAR"
    ws['A1'].font = title_font
    ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    # Thời gian xuất
    ws.merge_cells('A2:G2')
    ws['A2'] = f"Thời gian cập nhật: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    ws['A2'].font = meta_font
    ws['A2'].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 18

    # Dòng trống
    ws.row_dimensions[3].height = 10

    # Header các cột
    headers = [
        "STT", "Họ và Tên", "Số Điện Thoại", "Email",
        "Công Ty / Cửa Hàng", "Nhu Cầu / Nội Dung Tư Vấn", "Thời Gian Đăng Ký", "Trạng Thái"
    ]
    
    for col_num, header_title in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col_num)
        cell.value = header_title
        cell.font = header_font
        cell.fill = orange_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border
    ws.row_dimensions[4].height = 28

    # Đổ dữ liệu
    row_num = 5
    for idx, item in enumerate(queryset, 1):
        created_str = item.created_at.strftime('%d/%m/%Y %H:%M') if item.created_at else ""
        status_str = "Đã xử lý" if getattr(item, 'is_processed', False) else "Chờ liên hệ"

        row_data = [
            idx,
            item.full_name or "",
            item.phone or "",
            item.email or "—",
            item.company or "—",
            item.message or item.interest or "",
            created_str,
            status_str
        ]

        for col_num, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = val
            cell.font = data_font
            cell.border = thin_border
            
            # Căn chỉnh
            if col_num in [1, 3, 7, 8]:
                cell.alignment = Alignment(horizontal="center", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

            # Màu trạng thái
            if col_num == 8:
                if status_str == "Đã xử lý":
                    cell.font = Font(name="Arial", size=10, bold=True, color="2E7D32")
                else:
                    cell.font = Font(name="Arial", size=10, bold=True, color="D84315")

        ws.row_dimensions[row_num].height = 24
        row_num += 1

    # Tự động căn chỉnh độ rộng cột
    col_widths = {
        'A': 8,   # STT
        'B': 24,  # Họ tên
        'C': 16,  # Số điện thoại
        'D': 26,  # Email
        'E': 26,  # Công ty
        'F': 45,  # Nội dung
        'G': 20,  # Thời gian
        'H': 15   # Trạng thái
    }
    for col_letter, width in col_widths.items():
        ws.column_dimensions[col_letter].width = width

    return wb


def export_consultation_to_docx(consultation_obj, doc=None):
    """
    Xuất một phiếu yêu cầu tư vấn chi tiết sang file Word (.docx)
    """
    if doc is None:
        doc = Document()

    # Cài đặt margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Tiêu đề Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CÔNG TY TNHH KEO VỮA SÔNG HỒNG - SHK MORTAR\n")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(13)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(230, 90, 30) # Orange brand

    sub_run = title_p.add_run("PHIẾU TIẾP NHẬN THÔNG TIN ĐĂNG KÝ TƯ VẤN KHÁCH HÀNG\n")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(15)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(17, 17, 17)

    created_str = consultation_obj.created_at.strftime('%d/%m/%Y %H:%M:%S') if consultation_obj.created_at else datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    date_run = title_p.add_run(f"(Mã tiếp nhận: #SHK-{consultation_obj.pk:05d} — Ngày ghi nhận: {created_str})\n")
    date_run.font.name = "Arial"
    date_run.font.size = Pt(9.5)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    # Bảng thông tin khách hàng
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    fields = [
        ("Họ và tên khách hàng:", consultation_obj.full_name or "—"),
        ("Số điện thoại liên hệ:", consultation_obj.phone or "—"),
        ("Địa chỉ Email:", consultation_obj.email or "—"),
        ("Tên công ty / Cửa hàng:", consultation_obj.company or "—"),
        ("Mục quan tâm / Danh mục:", consultation_obj.interest or "Tư vấn sản phẩm"),
        ("Nội dung chi tiết / Địa chỉ:", consultation_obj.message or "—"),
        ("Trạng thái xử lý:", "Đã xử lý" if getattr(consultation_obj, 'is_processed', False) else "Chờ liên hệ tư vấn")
    ]

    for i, (label, val) in enumerate(fields):
        row = table.rows[i]
        
        # Label cell
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.2)
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.name = "Arial"
        r_lbl.font.size = Pt(10.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(30, 30, 30)

        # Value cell
        cell_val = row.cells[1]
        cell_val.width = Inches(4.5)
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(str(val))
        r_val.font.name = "Arial"
        r_val.font.size = Pt(10.5)
        r_val.font.color.rgb = RGBColor(60, 60, 60)

    # Ghi chú ký tên
    doc.add_paragraph("\n")
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    row_0 = sig_table.rows[0]
    p_s1 = row_0.cells[0].paragraphs[0]
    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s1 = p_s1.add_run("NGƯỜI TIẾP NHẬN\n(Ký và ghi rõ họ tên)")
    r_s1.font.name = "Arial"
    r_s1.font.size = Pt(10)
    r_s1.font.bold = True

    p_s2 = row_0.cells[1].paragraphs[0]
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s2 = p_s2.add_run("TRƯỞNG PHÒNG KINH DOANH\n(Ký duyệt)")
    r_s2.font.name = "Arial"
    r_s2.font.size = Pt(10)
    r_s2.font.bold = True

    return doc


def auto_save_consultation_to_files(consultation_obj):
    """
    Tự động lưu/cập nhật vào file Excel tổng hợp và xuất file Word cho từng lượt đăng ký
    """
    export_dir = get_export_dir()

    # 1. Cập nhật vào file Excel tổng hợp
    excel_path = os.path.join(export_dir, 'Danh_Sach_Dang_Ky_Tu_Van_SHK.xlsx')
    from .models import ConsultationRequest
    all_consultations = ConsultationRequest.objects.all().order_by('-created_at')
    wb = export_consultations_to_excel(all_consultations)
    wb.save(excel_path)

    # 2. Lưu file Word phiếu tư vấn cho khách hàng này
    safe_phone = str(consultation_obj.phone).replace(' ', '')
    docx_filename = f"Phieu_Tu_Van_SHK_{consultation_obj.pk}_{safe_phone}.docx"
    docx_path = os.path.join(export_dir, docx_filename)
    doc = export_consultation_to_docx(consultation_obj)
    doc.save(docx_path)

    return excel_path, docx_path
